from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_PATH = OUT_DIR / "Runk_Portfolio_Report.docx"

INK = RGBColor(25, 28, 34)
MUTED = RGBColor(93, 101, 115)
ACCENT = RGBColor(28, 87, 151)
GREEN = RGBColor(28, 124, 88)
LIGHT_BLUE = "EAF2FF"
LIGHT_GREEN = "EAF7F1"
LIGHT_GRAY = "F4F6F8"
WHITE = "FFFFFF"


def font(run, size: int = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    if fill:
        shade(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    font(run, 9, bold, INK)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    table_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if table_borders is None:
        table_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(table_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = table_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            table_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DDE2E8")


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    borders(t)

    for idx, header in enumerate(headers):
        set_cell(t.rows[0].cells[idx], header, True, LIGHT_BLUE)
        t.rows[0].cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value)
            if len(value) <= 8:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if widths:
        for row in t.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)

    doc.add_paragraph()
    return t


def callout(doc: Document, title: str, body: str, fill: str = LIGHT_GREEN) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t)
    cell = t.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    font(r, 10, True, ACCENT)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    font(r2, 9, False, INK)
    doc.add_paragraph()


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        font(run, 16 if level == 1 else 12, True, ACCENT if level == 1 else GREEN)


def para(doc: Document, text: str = "", size: int = 10, color: RGBColor = INK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    font(r, size, False, color)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        font(r, 10, False, INK)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("RunK Portfolio Report")
    font(r, 8, False, MUTED)


def cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(84)
    r = p.add_run("RunK\n포트폴리오 보고서")
    font(r, 28, True, INK)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(14)
    r2 = sub.add_run("사람과 사람을 잇는 소셜 러닝 MVP")
    font(r2, 14, False, MUTED)

    doc.add_paragraph()
    table(
        doc,
        ["항목", "내용"],
        [
            ["프로젝트", "RunK"],
            ["유형", "1인 개발 포트폴리오 MVP"],
            ["기술", "Flutter, FastAPI, MySQL, SQLAlchemy, JWT, Docker"],
            ["상태", "인증, 기록, 피드, 화면 구조, 테마 설정 구현"],
        ],
        [4, 11],
    )
    doc.add_page_break()


def build_doc() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    configure(doc)
    cover(doc)

    heading(doc, "1. 프로젝트 개요")
    para(
        doc,
        "RunK는 러닝 기록을 저장하고 친구와 공유하는 한국어 소셜 러닝 앱 MVP입니다. "
        "실무형 포트폴리오를 목표로 인증, API, DB, 앱 화면 구조, 보안 고려사항, 테스트 문서를 함께 정리했습니다.",
    )
    table(
        doc,
        ["구분", "내용"],
        [
            ["핵심 가치", "사람과 사람을 잇는 소셜 러닝"],
            ["대상 사용자", "러닝 기록을 남기고 친구와 공유하고 싶은 사용자"],
            ["MVP 범위", "회원가입, 로그인, 기록 저장, 피드 조회, 앱 주요 화면"],
            ["다음 단계", "친구 관계 API, 채팅, GPS 경로 저장, 소셜 로그인"],
        ],
        [4, 11],
    )

    heading(doc, "2. 주요 기능")
    table(
        doc,
        ["영역", "구현 내용"],
        [
            ["인증", "회원가입, 로그인, JWT 발급, 세션 복원, 이메일/닉네임 중복 체크"],
            ["보안", "bcrypt 비밀번호 해시, Bearer 토큰, 환경변수 기반 설정, .env 제외"],
            ["러닝", "거리, 시간, 날짜, 메모 저장 및 페이스 자동 계산"],
            ["소셜", "피드 화면, 친구 화면 UI, 채팅 탭 UI"],
            ["앱 경험", "스플래시, 홈, 기록, 설정, 라이트/다크 테마"],
        ],
        [4, 11],
    )

    heading(doc, "3. 시스템 아키텍처")
    callout(doc, "Architecture", "Flutter App -> REST API -> FastAPI -> SQLAlchemy -> MySQL", LIGHT_BLUE)
    table(
        doc,
        ["구성 요소", "역할"],
        [
            ["Flutter", "사용자 화면, 입력 검증, REST API 호출, 로컬 세션 저장"],
            ["FastAPI", "인증, 사용자, 러닝 기록, 피드 API 제공"],
            ["MySQL", "사용자와 러닝 기록 저장"],
            ["Docker Compose", "개발용 MySQL 실행 환경"],
            ["PowerShell Scripts", "DB, 백엔드, smoke test 실행 자동화"],
        ],
        [4, 11],
    )

    heading(doc, "4. API 명세 요약")
    table(
        doc,
        ["Method", "Endpoint", "설명", "인증"],
        [
            ["GET", "/health", "서버 상태 확인", "불필요"],
            ["GET", "/auth/check-email", "이메일 중복 확인", "불필요"],
            ["GET", "/auth/check-username", "닉네임 중복 확인", "불필요"],
            ["POST", "/auth/signup", "회원가입", "불필요"],
            ["POST", "/auth/login", "로그인", "불필요"],
            ["GET", "/users/me", "내 프로필 조회", "필요"],
            ["POST", "/running-records", "러닝 기록 저장", "필요"],
            ["GET", "/running-records/me", "내 러닝 기록 조회", "필요"],
            ["GET", "/feed", "피드 조회", "불필요"],
        ],
        [2.2, 4.8, 5.5, 2.0],
    )

    heading(doc, "5. DB 설계")
    table(
        doc,
        ["테이블", "핵심 컬럼", "설명"],
        [
            ["users", "id, email, username, hashed_password, created_at", "회원 계정과 인증 정보를 저장"],
            ["running_records", "id, user_id, distance_km, duration_seconds, pace, run_date, memo", "사용자별 러닝 기록 저장"],
        ],
        [3.2, 6.5, 5.3],
    )

    heading(doc, "6. 테스트 및 검증")
    table(
        doc,
        ["항목", "명령/방법", "기대 결과"],
        [
            ["정적 분석", "flutter analyze", "Dart/Flutter 분석 통과"],
            ["위젯 테스트", "flutter test", "테스트 통과"],
            ["Windows 빌드", "flutter build windows --debug", "데스크톱 빌드 성공"],
            ["백엔드 smoke", "scripts/smoke-backend.ps1", "회원가입, 로그인, 기록, 피드 흐름 확인"],
        ],
        [3.2, 5.7, 6.1],
    )

    heading(doc, "7. 포트폴리오 어필 포인트")
    bullets(
        doc,
        [
            "단순 화면 구현이 아니라 백엔드, DB, 인증, 보안, 테스트 문서까지 연결했습니다.",
            "회원가입 입력 검증과 중복 체크를 프론트와 API 양쪽에서 고려했습니다.",
            "실제 유지보수를 위해 화면, 서비스, 모델, 테마, 위젯을 분리했습니다.",
            "README, 아키텍처, API, DB, 테스트 계획을 제출 가능한 형태로 정리했습니다.",
        ],
    )

    heading(doc, "8. 다음 개발 계획")
    table(
        doc,
        ["우선순위", "작업"],
        [
            ["1", "친구 관계 API와 친구 목록 실데이터 연동"],
            ["2", "러닝 기록 작성 화면과 피드 업로드 흐름 고도화"],
            ["3", "GPS 위치 권한, 경로 저장, 지도 표시"],
            ["4", "소셜 로그인 도입 전 OAuth 보안 검토"],
            ["5", "Android 실기기 QA와 배포 준비"],
        ],
        [3, 12],
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
